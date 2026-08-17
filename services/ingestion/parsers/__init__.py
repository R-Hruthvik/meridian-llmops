"""Multi-format document parsers for PDF, Markdown, HTML, DOCX, and TXT."""

import io
import uuid

from bs4 import BeautifulSoup

from packages.core.models import Document, DocumentFormat


def parse_markdown(content: str) -> str:
    # Retain Markdown structure
    return content.strip()


def parse_html(content: str) -> str:
    soup = BeautifulSoup(content, "html.parser")
    # Extract clean text with heading preservation
    for heading in soup.find_all(["h1", "h2", "h3", "h4"]):
        level = heading.name[1]
        heading.replace_with(f"\n\n{'#' * int(level)} {heading.get_text().strip()}\n")
    return soup.get_text().strip()


def parse_pdf(content_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content_bytes))
    pages_text = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages_text.append(f"## Page {i + 1}\n{text}")
    return "\n\n".join(pages_text).strip()


def parse_docx(content_bytes: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(content_bytes))
    paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip():
            if p.style and p.style.name.startswith("Heading"):
                paragraphs.append(f"## {p.text.strip()}")
            else:
                paragraphs.append(p.text.strip())
    return "\n\n".join(paragraphs).strip()


def parse_document(
    content: str | bytes,
    title: str,
    doc_format: DocumentFormat,
    source: str,
    doc_id: str | None = None,
) -> Document:
    """Parses raw content into a structured Document domain model."""
    doc_id = doc_id or f"doc-{uuid.uuid4().hex[:8]}"

    if doc_format == DocumentFormat.PDF and isinstance(content, bytes):
        text = parse_pdf(content)
    elif doc_format == DocumentFormat.DOCX and isinstance(content, bytes):
        text = parse_docx(content)
    elif doc_format == DocumentFormat.HTML:
        text = parse_html(content if isinstance(content, str) else content.decode("utf-8", errors="ignore"))
    elif doc_format == DocumentFormat.MARKDOWN:
        text = parse_markdown(content if isinstance(content, str) else content.decode("utf-8", errors="ignore"))
    else:
        text = content if isinstance(content, str) else content.decode("utf-8", errors="ignore")

    return Document(
        id=doc_id,
        title=title,
        text=text,
        format=doc_format,
        source=source,
    )
